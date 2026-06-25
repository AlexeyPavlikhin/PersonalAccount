#!/usr/bin/env python3
"""
Сборка демонстрационного DOCX-шаблона договора с плейсхолдерами ${field_code}.

Имена полей = плейсхолдеры ${field_code}; field_map_json в БД — {} (особый маппинг не требуется).
Плейсхолдеры заполняются на сервере (inc/docx_form_filler.php).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATHS = [
    ROOT / "requirements" / "docs_temlpates" / "demo_legal_services_agreement.docx",
    ROOT / "requirements" / "docs_temlpates" / "doc" / "demo_legal_services_agreement.docx",
]


def add_line(doc: Document, parts: list[tuple[str, str]]) -> None:
    """Абзац из чередующихся статичного текста и плейсхолдеров ${name}."""
    paragraph = doc.add_paragraph()
    for kind, value in parts:
        run = paragraph.add_run(value if kind == "text" else f"${{{value}}}")
        run.font.size = Pt(11)


def build_docx(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(15)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Договор оказания юридических услуг (демо)")
    run.bold = True
    run.font.size = Pt(14)

    add_line(doc, [("text", "г. Москва     «"), ("field", "contract_day"), ("text", "» "), ("field", "contract_month_num"), ("text", " "), ("field", "contract_month_words"), ("text", " 20"), ("field", "contract_year_yy"), ("text", " / "), ("field", "contract_year_yyyy"), ("text", " г.")])
    doc.add_paragraph()

    add_line(doc, [("text", "Наименование (краткое ОПФ): "), ("field", "company_name_short_opf")])
    add_line(doc, [("text", "Наименование (полное ОПФ): "), ("field", "company_name_full_opf")])
    add_line(doc, [("text", "Наименование (основное): "), ("field", "company_name")])
    add_line(doc, [("text", "ОПФ (кратк.): "), ("field", "opf_short"), ("text", "    ОПФ (полн.): "), ("field", "opf_full")])
    doc.add_paragraph()

    add_line(doc, [("text", "В лице "), ("field", "signer_position_genitive"), ("text", " "), ("field", "signer_name_genitive"), ("text", ", действующего на основании "), ("field", "signer_basis"), ("text", ", с одной стороны, и Исполнитель — с другой, заключили настоящий договор.")])
    doc.add_paragraph()

    add_line(doc, [("text", "Форма обращения к контрагенту: "), ("field", "party_named_form"), ("text", " в дальнейшем «Заказчик».")])
    doc.add_paragraph()

    add_line(doc, [("text", "Подписант (именительный падеж): "), ("field", "signer_name")])
    add_line(doc, [("text", "Должность (именительный падеж): "), ("field", "signer_position")])
    add_line(doc, [("text", "Подпись в реквизитах: "), ("field", "signer_short"), ("text", " / "), ("field", "signer_initials")])
    doc.add_paragraph()

    add_line(doc, [("text", "Адрес: "), ("field", "address")])
    add_line(doc, [("text", "ОГРН: "), ("field", "ogrn"), ("text", "    ИНН: "), ("field", "inn"), ("text", "    КПП: "), ("field", "kpp")])
    add_line(doc, [("text", "Р/с: "), ("field", "checking_account"), ("text", "    Банк: "), ("field", "bank_name")])
    add_line(doc, [("text", "К/с: "), ("field", "corr_account"), ("text", "    БИК: "), ("field", "bik")])
    add_line(doc, [("text", "E-mail: "), ("field", "email"), ("text", "    Телефон: "), ("field", "phone")])
    doc.add_paragraph()

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Демо: наименование с кратк./полн. ОПФ, дата, склонения, инициалы, «именуемый/именуемая»."
    )
    note_run.font.size = Pt(9)

    doc.save(str(output_path))


def main() -> None:
    for output_path in OUTPUT_PATHS:
        build_docx(output_path)
        print(f"Written: {output_path}")


if __name__ == "__main__":
    main()
