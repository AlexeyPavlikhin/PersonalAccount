#!/usr/bin/env python3
"""
Сборка тестового DOCX: таблица «имя поля — значение» для всех поддерживаемых field_code.

Плейсхолдеры ${field_code}; field_map_json шаблона — {} (базовая подстановка всех ключей form_data/enrich
передаются в DOCX как есть). См. database/load_document_template_fields_debug.sql.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATHS = [
    ROOT / "requirements" / "docs_temlpates" / "all_fields_debug.docx",
    ROOT / "requirements" / "docs_temlpates" / "doc" / "all_fields_debug.docx",
]

# Полный перечень field_code (форма, enrich, API, реестр, служебные)
ALL_FIELD_CODES: list[tuple[str, str]] = [
    ("contract_number", "Номер договора (XX-MM/YY)"),
    ("contract_subject_short", "Краткое описание предмета договора"),
    ("contract_date", "Дата составления (ГГГГ-ММ-ДД)"),
    ("contract_date_display", "Дата для колонтитула (ДД.ММ.ГГГГ)"),
    ("contract_day", "День договора"),
    ("contract_month_num", "Месяц (число, 01–12)"),
    ("contract_month_words", "Месяц прописью (род. падеж)"),
    ("contract_year_yyyy", "Год (YYYY)"),
    ("contract_year_yy", "Год (YY)"),
    ("spec_number", "Номер спецификации"),
    ("spec_date", "Дата спецификации"),
    ("invoice_number", "Номер счёта"),
    ("invoice_date", "Дата счёта"),
    ("planned_act_date", "Плановая дата акта (spec_date + 14 дней)"),
    ("counterparty_display", "Контрагент (снимок для реестра)"),
    ("company_name_short_opf", "Наименование (краткое ОПФ)"),
    ("company_name_full_opf", "Наименование (полное ОПФ)"),
    ("company_name", "Наименование (основное)"),
    ("opf_short", "ОПФ (краткое)"),
    ("opf_full", "ОПФ (полное)"),
    ("signer_name", "ФИО подписанта (именительный)"),
    ("signer_name_genitive", "ФИО подписанта (родительный)"),
    ("signer_position", "Должность (именительный)"),
    ("signer_position_genitive", "Должность (родительный)"),
    ("signer_basis", "Основание полномочий"),
    ("signer_initials", "Инициалы"),
    ("signer_short", "Подпись «Фамилия И.О.»"),
    ("signer_signature", "Подпись в реквизитах (как signer_short)"),
    ("party_named_form", "именуемый / именуемая / именуемое"),
    ("address", "Адрес"),
    ("ogrn", "ОГРН / ОГРНИП"),
    ("inn", "ИНН"),
    ("kpp", "КПП"),
    ("inn_kpp", "ИНН/КПП одной строкой"),
    ("checking_account", "Расчётный счёт"),
    ("bank_name", "Банк"),
    ("corr_account", "Корреспондентский счёт"),
    ("bik", "БИК"),
    ("swift", "SWIFT банка"),
    ("email", "E-mail"),
    ("phone", "Телефон"),
    ("counterparty_type", "Тип контрагента (DaData)"),
    ("counterparty_is_individual", "Физлицо (1/0)"),
    ("signer_gender", "Пол подписанта (MALE/FEMALE/UNKNOWN)"),
]

DEMO_TABLE_COLUMNS = ("demo_service_name", "demo_service_price")


def build_docx(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Справочник полей шаблона (тест)")
    run.bold = True
    run.font.size = Pt(14)

    subtitle = doc.add_paragraph(
        "Таблица для проверки подстановки: в колонке «Значение» — результат enrich и формы."
    )
    subtitle.runs[0].font.size = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Имя поля (field_code)"
    header_cells[1].text = "Значение"
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for r in paragraph.runs:
                r.bold = True
                r.font.size = Pt(10)

    for field_code, _label in ALL_FIELD_CODES:
        row = table.add_row().cells
        name_cell = row[0]
        value_cell = row[1]
        name_para = name_cell.paragraphs[0]
        name_run = name_para.add_run(f"{field_code} — ")
        name_run.font.size = Pt(9)
        name_run.bold = True
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(f"${{{field_code}}}")
        value_run.font.size = Pt(9)

    doc.add_paragraph()
    demo_title = doc.add_paragraph("Демо табличного блока (demo_services)")
    demo_title.runs[0].bold = True
    demo_title.runs[0].font.size = Pt(11)

    demo_table = doc.add_table(rows=2, cols=2)
    demo_table.style = "Table Grid"
    demo_header = demo_table.rows[0].cells
    demo_header[0].text = "Услуга"
    demo_header[1].text = "Цена"
    demo_row = demo_table.rows[1].cells
    demo_row[0].text = "${#demo_services} ${demo_service_name}"
    demo_row[1].text = f"${{{DEMO_TABLE_COLUMNS[1]}}}"

    doc.save(str(output_path))


def main() -> None:
    for output_path in OUTPUT_PATHS:
        build_docx(output_path)
        print(f"Written: {output_path}")


if __name__ == "__main__":
    main()
